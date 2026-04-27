from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
TEZ_DIR = ROOT / "README" / "Tez"
TEMPLATE = TEZ_DIR / "BOP_Akademik_Bitirme_Ara_Taslak_APA7.docx"
OUTPUT = TEZ_DIR / "MES_Siber_Fiziksel_Konveyor_Ara_Raporu_APA7.docx"
ASSET_DIR = TEZ_DIR / "_mes_ara_rapor_assets"


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_in: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for i, width in enumerate(widths):
        if i < len(table.columns):
            table.columns[i].width = Inches(width)
        if grid is not None and i < len(grid.gridCol_lst):
            grid.gridCol_lst[i].set(qn("w:w"), str(int(width * 1440)))


def format_run(run, size=11, bold=False, color=None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, align=None, before=0, after=6, line_spacing=1.08) -> None:
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    for run in paragraph.runs:
        if paragraph.style.name in {"Title"}:
            format_run(run, size=17, bold=True, color="1F4E79")
        elif paragraph.style.name in {"Subtitle"}:
            format_run(run, size=11, bold=False, color="404040")
        elif paragraph.style.name in {"Heading 1"}:
            format_run(run, size=14, bold=True, color="1F4E79")
        elif paragraph.style.name in {"Heading 2"}:
            format_run(run, size=12, bold=True, color="2F5597")
        elif paragraph.style.name in {"Reference"}:
            format_run(run, size=10)
        else:
            format_run(run, size=11)


def add_p(doc: Document, text: str = "", style: str = "Body Text", align=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    format_paragraph(p, align=align, before=before, after=after)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    format_paragraph(p, before=10 if level == 1 else 6, after=5)
    return p


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    title_para = add_heading(doc, title, 2)
    title_para.paragraph_format.keep_with_next = False
    image_path = render_table_image(title, headers, rows, widths)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_after = Pt(8)
    p.add_run().add_picture(str(image_path), width=Inches(6.25))


def load_table_font(size: int, bold: bool = False):
    font_candidates = [
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for candidate in font_candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = str(text).split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def draw_multiline_centered(draw, box, lines, font, fill, line_gap=5, center=False):
    x1, y1, x2, y2 = box
    line_h = draw.textbbox((0, 0), "Ag", font=font)[3] + line_gap
    total_h = line_h * len(lines) - line_gap
    y = y1 + max(0, (y2 - y1 - total_h) // 2)
    for line in lines:
        text_w = draw.textbbox((0, 0), line, font=font)[2]
        x = x1 + ((x2 - x1 - text_w) // 2 if center else 0)
        draw.text((x, y), line, fill=fill, font=font)
        y += line_h


def render_table_image(title: str, headers: list[str], rows: list[list[str]], widths: list[float]) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    idx = len(list(ASSET_DIR.glob("table_*.png"))) + 1
    path = ASSET_DIR / f"table_{idx:02d}.png"

    canvas_w = 1500
    margin = 34
    table_w = canvas_w - (2 * margin)
    total_width = sum(widths)
    col_w = [max(120, int(table_w * w / total_width)) for w in widths]
    col_w[-1] += table_w - sum(col_w)

    header_font = load_table_font(29, bold=True)
    body_font = load_table_font(27, bold=False)
    probe = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(probe)
    pad_x = 18
    pad_y = 14
    body_line_h = draw.textbbox((0, 0), "Ag", font=body_font)[3] + 7
    header_line_h = draw.textbbox((0, 0), "Ag", font=header_font)[3] + 7

    wrapped_header = [
        wrap_text(draw, h, header_font, col_w[i] - (2 * pad_x)) for i, h in enumerate(headers)
    ]
    row_wraps = []
    for row in rows:
        row_wraps.append(
            [wrap_text(draw, row[i], body_font, col_w[i] - (2 * pad_x)) for i in range(len(headers))]
        )

    header_h = max(len(lines) * header_line_h for lines in wrapped_header) + 2 * pad_y
    row_heights = [
        max(len(cell_lines) * body_line_h for cell_lines in wrapped_row) + 2 * pad_y
        for wrapped_row in row_wraps
    ]
    canvas_h = margin + header_h + sum(row_heights) + margin

    image = Image.new("RGB", (canvas_w, canvas_h), "white")
    draw = ImageDraw.Draw(image)
    border = "#6B7A8C"
    header_fill = "#D9EAF7"
    text = "#1F1F1F"
    header_text = "#1F4E79"

    x = margin
    y = margin
    for i, lines in enumerate(wrapped_header):
        draw.rectangle([x, y, x + col_w[i], y + header_h], fill=header_fill, outline=border, width=2)
        draw_multiline_centered(
            draw,
            (x + pad_x, y + pad_y, x + col_w[i] - pad_x, y + header_h - pad_y),
            lines,
            header_font,
            header_text,
            center=True,
        )
        x += col_w[i]
    y += header_h

    for wrapped_row, row_h in zip(row_wraps, row_heights):
        x = margin
        for i, lines in enumerate(wrapped_row):
            draw.rectangle([x, y, x + col_w[i], y + row_h], fill="white", outline=border, width=2)
            draw_multiline_centered(
                draw,
                (x + pad_x, y + pad_y, x + col_w[i] - pad_x, y + row_h - pad_y),
                lines,
                body_font,
                text,
                center=False,
            )
            x += col_w[i]
        y += row_h

    image.save(path)
    return path


def add_bullet(doc: Document, text: str) -> None:
    add_p(doc, f"• {text}", "Body Text", after=3)


def add_references(doc: Document, refs: list[str]) -> None:
    for ref in refs:
        p = doc.add_paragraph(style="Reference")
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(ref)
        format_run(run, size=10)


def build() -> None:
    doc = Document(TEMPLATE)
    clear_document(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    for name in ["Normal", "Body Text", "Reference"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(11 if name != "Reference" else 10)
    for name in ["Title", "Subtitle", "Heading 1", "Heading 2"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    add_p(doc, "MES TABANLI SİBER-FİZİKSEL KONVEYÖR SİSTEMİ ARA RAPOR TASLAĞI", "Title", WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_p(doc, "Endüstri Mühendisliği Bakışıyla MES, OEE, İzlenebilirlik ve Düşük Maliyetli Dijital Üretim Mimarisi", "Subtitle", WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_p(doc, "Ara Rapor - APA 7 Atıf Düzeni", "Subtitle", WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_p(doc, "Ana proje kaynağı: MES repo dokümanları, Arduino Mega/ESP32 gömülü kodları ve mes_web uygulama katmanı.", "Normal", WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_p(doc, "Deneysel sistem kaynağı: Fiziksel mini konveyör, robot kol, TCS3200 renk algılama, MQTT veri akışı ve Excel tabanlı MES kayıt yapısı.", "Normal", WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_p(doc, "Hazırlayan: [Doldurulacak]", "Normal", WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_p(doc, "Danışman: [Doldurulacak]", "Normal", WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_p(doc, "Tarih: 25 Nisan 2026", "Normal", WD_ALIGN_PARAGRAPH.CENTER, after=14)

    add_heading(doc, "Özet", 1)
    add_p(doc, "Bu ara rapor, laboratuvar ölçekli renk ayırma konveyör sisteminin yalnızca çalışan bir otomasyon düzeneği olarak değil, siber-fiziksel üretim sistemi ve üretim yürütme sistemi (MES) prototipi olarak ele alınabileceğini göstermek üzere hazırlanmıştır. Projede fiziksel katmanda konveyör, robot kol, renk sensörü ve limit switch yapısı; kontrol katmanında Arduino Mega; köprü katmanında ESP32; haberleşme katmanında MQTT; uygulama katmanında ise FastAPI, WebSocket, operator kiosk, teknisyen ekranı, OEE runtime ve Excel tabanlı veri tabanı bulunmaktadır. Bu katmanlı yapı, Endüstri 4.0 literatüründe vurgulanan fiziksel sistem ile siber karar katmanı arasındaki gerçek zamanlı veri bağlantısını küçük ölçekli ve doğrulanabilir bir üretim hattı üzerinde somutlaştırmaktadır.", after=6)
    add_p(doc, "Endüstri mühendisliği açısından çalışmanın ana değeri, üretim olaylarının ölçülebilir ve yorumlanabilir hale getirilmesidir. Renk sensörüyle sınıflandırılan ürünler kuyruk mantığına alınmakta, robot kol ile alma-bırakma çevrimi tamamlanmakta, üretim olayları MQTT üzerinden MES katmanına taşınmakta ve tamamlanan ürün, kalite, arıza, bakım, iş emri ve OEE bilgileri günlük workbook içine yazılmaktadır. Böylece sistem; çevrim süresi, plansız duruş, planlı bakım, kalite düzeltme, stok davranışı ve operatör aksiyonları gibi klasik endüstri mühendisliği karar alanlarını veriyle ilişkilendiren bir eğitim ve araştırma platformu sunmaktadır.", after=6)
    add_p(doc, "Akademik çerçevede çalışma; MES ve ISA-95, siber-fiziksel üretim sistemleri, MQTT tabanlı IoT haberleşmesi, OEE/KPI yönetimi ve dijital ikiz yaklaşımları ile ilişkilendirilmektedir. ISA-95 standardı MES/MOM katmanını işletme planlama sistemleri ile saha kontrol sistemleri arasında konumlandırırken, Lee, Bagheri ve Kao (2015) siber-fiziksel üretim sistemlerinde fiziksel varlıklar ile siber hesaplama alanı arasındaki senkronizasyonu vurgulamaktadır. Bu proje, bu iki yaklaşımı düşük maliyetli bir konveyör prototipi üzerinde birleştirerek izlenebilir, ölçülebilir ve geliştirilebilir bir MES ara raporu ortaya koymaktadır.", after=8)

    add_heading(doc, "Anahtar Kelimeler", 2)
    add_p(doc, "MES, siber-fiziksel sistem, Endüstri 4.0, MQTT, Arduino Mega, ESP32, OEE, üretim izlenebilirliği, robot kol, konveyör, Excel tabanlı veri tabanı, operator kiosk, teknisyen çağrı sistemi.", after=8)

    add_heading(doc, "1. Problemin Tanımı", 1)
    add_p(doc, "Günümüz üretim sistemlerinde temel problem yalnızca makinenin çalışması değil, makinede gerçekleşen olayların işletme kararlarına dönüşebilecek anlamlı veri halinde yakalanmasıdır. Fiziksel sahada bir sensörün ürün algılaması, bir robot kolun ürünü taşıması veya bir operatörün arıza bildirmesi tek başına yeterli değildir; bu olayların zaman, ürün kimliği, kalite durumu, iş emri ve performans göstergeleriyle ilişkilendirilmesi gerekir. MES sistemleri bu nedenle ERP ve üretim planlama katmanı ile saha otomasyonu arasında kritik bir köprü görevi görür. ISA-95 standardının işletme ve kontrol sistemleri arasındaki veri alışverişini tanımlaması bu ihtiyacı kurumsal düzeyde ortaya koymaktadır (International Society of Automation, n.d.).", after=6)
    add_p(doc, "Üniversite laboratuvarlarında geliştirilen birçok otomasyon prototipi fiziksel hareketleri başarıyla gerçekleştirse de üretim yürütme, izlenebilirlik ve performans yönetimi açısından eksik kalmaktadır. Bir konveyör hattı renkli parçaları algılayıp ayırabilir; ancak hangi ürünün ne zaman algılandığı, hangi kararla ayıklandığı, kaç ürünün iyi veya hurda olduğu, hangi duruşların OEE hesabına girdiği ve hangi verinin günlük rapora dönüştüğü izlenemiyorsa sistem endüstri mühendisliği açısından sınırlı değer üretir.", after=6)
    add_p(doc, "Bu projede çözülen temel problem, düşük maliyetli bir konveyör-robot kol düzeneğinin MES bakış açısıyla yeniden ele alınmasıdır. Sistem yalnızca ürünleri renklerine göre ayırmakla kalmamakta; olay logu, dashboard snapshot, kiosk aksiyonu, teknisyen çağrısı, iş emri, OEE runtime state ve Excel workbook kayıtlarıyla üretim verisini yapılandırmaktadır. Bu nedenle proje, otomasyon, bilişim ve endüstri mühendisliği kesişiminde yer alan siber-fiziksel bir üretim yürütme prototipi olarak değerlendirilmelidir.", after=6)
    add_p(doc, "Endüstri mühendisliği açısından problem; çevrim süresi, kalite kaybı, duruş sınıflandırması, iş emri takibi, bakım kaydı ve görsel yönetim eksenlerinde ele alınmaktadır. OEE literatüründe kullanılabilirlik, performans ve kalite bileşenleri üretim etkinliğini ölçmek için birlikte değerlendirilir (Ng Corrales et al., 2020). Mevcut proje bu mantığı laboratuvar ölçeğine indirerek, fiziksel hat olaylarını gerçek zamanlı dashboard ve kalıcı workbook kayıtlarıyla ilişkilendirmektedir.", after=8)

    add_heading(doc, "2. Problemin Amacı", 1)
    add_p(doc, "Bu çalışmanın genel amacı, mini konveyör hattının MES tabanlı siber-fiziksel üretim sistemi olarak tasarlanabilir ve raporlanabilir olduğunu göstermektir. Çalışma; fiziksel cihaz, gömülü kontrol, IoT haberleşme, web tabanlı MES arayüzü, OEE hesaplama, kalite/arıza/bakım kayıtları ve Excel tabanlı veri sürekliliğini tek bir ara rapor çerçevesinde birleştirmektedir.", after=6)
    add_p(doc, "Birinci amaç, fiziksel üretim hattındaki sensör, konveyör ve robot kol olaylarını standart bir veri akışına dönüştürmektir. Arduino Mega bu noktada ürün algılama, renk sınıflandırma, kuyruk oluşturma ve robot kol tetikleme kararlarının ana otoritesi olarak çalışır. ESP32 ise bu olayları MQTT üzerinden MES katmanına taşıyarak fiziksel sistem ile siber izleme katmanı arasında köprü kurar.", after=6)
    add_p(doc, "İkinci amaç, MES katmanının üretim sahası için neden gerekli olduğunu göstermektir. `mes_web` uygulaması dashboard, operator kiosk, teknisyen ekranı, REST API, WebSocket yayını, komut gönderimi, OEE runtime yönetimi ve workbook yazımını tek yerde toplayarak MES'in yalnızca bir ekran değil, üretim yürütme ve karar destek katmanı olduğunu ortaya koyar. Kletti (2007), MES'in süreç yeteneği ve gerçek zamanlı üretim yönetimi için önemini vurgularken; bu proje aynı fikri küçük ölçekli, uygulanabilir bir prototiple göstermektedir.", after=6)
    add_p(doc, "Üçüncü amaç, endüstri mühendisliği metriklerini veri mimarisine bağlamaktır. OEE hesabında açılış kontrol süresi OEE dışı, kapanış/bakım süresi planlı duruş, manuel arıza süresi ise plansız duruş olarak sınıflandırılmaktadır. Bu ayrım, kullanılabilirlik hesabının hatalı yorumlanmasını önler ve üretim kayıplarını doğru kategorilerde analiz etmeyi sağlar.", after=6)
    add_p(doc, "Dördüncü amaç, projenin ileride FERP, dijital ikiz veya çok istasyonlu üretim senaryolarına genişletilebilecek veri temelini hazırlamaktır. Günlük Excel workbook, olay logu, ölçümler, tamamlanan ürünler, OEE anlıkları, vision verisi, iş emirleri, depo stokları ve bakım kayıtları gibi sheet'ler aracılığıyla hem insan tarafından okunabilir hem de yazılımsal entegrasyona uygun bir veri sınırı oluşturmaktadır.", after=8)

    add_heading(doc, "3. Problemin Kapsamı", 1)
    add_p(doc, "Çalışmanın kapsamı, laboratuvar ölçekli bir renk ayırma konveyör hattının fiziksel kontrol ve MES entegrasyonunun incelenmesiyle sınırlıdır. Sistem tam ölçekli bir fabrika otomasyonunu temsil etme iddiasında değildir; ancak MES, OEE, izlenebilirlik ve siber-fiziksel sistem kavramlarını kontrollü bir üretim hattı üzerinde test etmek için yeterli bileşenleri içermektedir.", after=6)

    add_heading(doc, "3.1 Fiziksel Üretim Katmanı", 2)
    add_p(doc, "Fiziksel katmanda mini konveyör, DC motor, motor sürücü, TCS3200 renk sensörü, dört servo motorlu robot kol, step motorlu doğrusal hareket yapısı ve iki limit switch bulunmaktadır. Konveyör ürünleri algılama bölgesine taşır, renk sensörü ürünün sınıfını belirler, robot kol ise ürünün alma ve bırakma çevrimini gerçekleştirir. Limit switch'ler alma ve bırakma noktalarının referanslanmasını sağlayarak mekanik tekrarlanabilirliği artırır.", after=6)

    add_heading(doc, "3.2 Kontrol ve Gömülü Yazılım Katmanı", 2)
    add_p(doc, "Kontrol katmanında Arduino Mega tabanlı yazılım yer almaktadır. Mega, konveyör motorunu sürmekte, renk sensörü örneklerini değerlendirmekte, ürünleri kuyruk mantığına almakta ve robot kol hareket sırasını yönetmektedir. Robot kol çevrimi bloklamayan bir yapı ile çalıştırıldığı için step motor, servo motor ve haberleşme işlemleri aynı döngü içinde izlenebilmektedir. Sistem `start`, `stop`, `status`, `pickplace` ve kalibrasyon komutlarını kabul edecek biçimde tasarlanmıştır.", after=6)

    add_heading(doc, "3.3 Fiziksel Konveyör ve Robot Kol Sistemi", 2)
    add_p(doc, "Konveyörün mekanik yapısı, ürünün sensör bölgesine düzenli biçimde taşınmasını ve ölçüm sırasında kısa süreli durdurulmasını esas alır. Renk sensöründen gelen değerler tek örnekle değil, birden fazla örnek üzerinden medyan, oy çokluğu ve renk skoru mantığıyla değerlendirilir. Bu yaklaşım saha ışığı, ürün konumu ve sensör gürültüsü gibi değişkenlerin sınıflandırma kalitesini bozmasını azaltır.", after=6)
    add_p(doc, "Robot kol, alma pozisyonu, kavrama, kaldırma, taşıma, bırakma ve geri dönüş adımlarından oluşan bir çevrim izler. Alma noktasında LIM22, bırakma noktasında LIM23 referansı kullanılmaktadır. Gripper hareketleri kademeli yürütülerek ürünün sert darbeyle tutulması önlenir. Step motor hareketlerinde zaman aşımı ve limit kontrolü bulunduğu için mekanik sıkışma veya hedefe ulaşamama gibi durumlar loglanabilir. Bu yapı, endüstri mühendisliği açısından standart operasyon akışı, çevrim süresi ve arıza nedeni analizi için temel veri üretir.", after=6)
    add_p(doc, "Kalibrasyon süreci boş zemin, kırmızı, sarı ve mavi renk referanslarının sahaya göre ayarlanmasını içerir. Komut arayüzünde `cal x`, `cal r`/`cal k`, `cal y`/`cal s` ve `cal b`/`cal m` seçenekleri bulunmaktadır. Kalibrasyon verilerinin EEPROM üzerinde tutulması, sistem yeniden başlatıldığında önceki saha ayarlarının korunmasına katkı sağlar.", after=6)

    add_heading(doc, "3.4 Haberleşme ve MES Katmanı", 2)
    add_p(doc, "Haberleşme katmanı Mega-ESP32-MQTT zinciri üzerine kuruludur. Arduino Mega olay satırlarını seri haberleşmeyle ESP32'ye iletir. ESP32, bu satırları `sau/iot/mega/konveyor/` kökü altında MQTT topic'lerine yayınlar. MES tarafındaki `mes_web` uygulaması topic'leri dinler, veriyi ayrıştırır, dashboard snapshot'ı üretir ve kullanıcı arayüzlerine REST ile WebSocket üzerinden canlı veri sağlar. Tarayıcı tabanlı dashboard, kiosk ve teknisyen ekranları MQTT'ye doğrudan bağlanmaz; bu karar, saha haberleşmesini backend tarafında merkezileştirerek bakım ve güvenilirlik açısından daha kontrollü bir mimari oluşturur.", after=6)
    add_p(doc, "MQTT'nin publish/subscribe modeli, düşük kaynaklı cihazların ve uygulama katmanının gevşek bağlı çalışmasını sağlar. OASIS MQTT standardı istemci-sunucu üzerinden topic temelli mesajlaşma yapısını tanımlamaktadır (OASIS, 2014). Projede ESP32, Mega'dan gelen `status`, `logs`, `heartbeat` ve `bridge/status` mesajlarını yayınlarken; `cmd` topic'i üzerinden gelen komutları tekrar Mega'ya aktarır. Böylece saha cihazları ile MES arayüzü arasında çift yönlü ama denetlenebilir bir veri yolu oluşur.", after=6)

    add_heading(doc, "3.5 Veri, OEE ve İzlenebilirlik Katmanı", 2)
    add_p(doc, "MES veri katmanı günlük Excel workbook, OEE runtime JSON dosyası ve canlı snapshot modellerinden oluşur. Workbook içinde `1_Olay_Logu`, `2_Olcumler`, `3_Arizalar`, `4_Uretim_Tamamlanan`, `5_OEE_Anliklari`, `6_Vision`, `7_Is_Emirleri`, `8_Depo_Stok`, `9_Bakim_Kayitlari` ve `99_Raw_Logs` sheet'leri bulunmaktadır. Bu yapı, fiziksel hattan gelen ham logların hem denetlenebilir biçimde saklanmasına hem de üretim yönetimi açısından yorumlanabilir tablolara dönüştürülmesine imkan verir.", after=6)
    add_p(doc, "OEE tarafında sistem aktif vardiya mantığına göre çalışmaktadır. Aktif vardiya başlamadan üretim sayımı ve performans değerlendirmesi başlatılmaz. Tamamlanan ürünler varsayılan olarak iyi kalite kabul edilir; ancak dashboard veya kiosk üzerinden `GOOD`, `REWORK` ve `SCRAP` kalite düzeltmesi yapılabilir. Hurda olarak işaretlenen ürünlerin depo stokuna düşmemesi, üretim ve stok tutarlılığı açısından önemli bir kontrol noktasıdır.", after=8)

    doc.add_page_break()
    add_table(
        doc,
        "Tablo 1. MES-CPS mimari katman haritası",
        ["Katman", "Proje karşılığı", "Ana görev", "Endüstri mühendisliği yorumu"],
        [
            ["Fiziksel", "Konveyör, robot kol, sensör, switch", "Ürünü algılama, taşıma ve ayırma", "Operasyon akışı ve standart çevrim gözlemi"],
            ["Kontrol", "Arduino Mega", "Motor, sensör, kuyruk ve robot kararları", "Saha olaylarının birincil karar otoritesi"],
            ["Edge/Bridge", "ESP32", "Seri veriyi MQTT'ye aktarma ve komut taşıma", "Fiziksel-siber veri bağlantısı"],
            ["Haberleşme", "MQTT broker", "Topic tabanlı mesajlaşma", "Gerçek zamanlı izleme ve ayrık sistem entegrasyonu"],
            ["MES", "mes_web", "Dashboard, kiosk, OEE, workbook, REST/WS", "Üretim yürütme, izlenebilirlik ve karar destek"],
            ["Gözlem", "Raspberry vision", "Pasif renk ve crossing kontrolü", "Kalite doğrulama ve çapraz kontrol potansiyeli"],
            ["Veri", "Excel workbook ve runtime JSON", "Kalıcı kayıt ve anlık durum", "Raporlama, KPI ve entegrasyon temeli"],
        ],
        [1.0, 1.55, 1.75, 1.95],
    )

    add_table(
        doc,
        "Tablo 2. MQTT topic ağacı ve veri akışı özeti",
        ["Topic", "Yayıncı", "Abone", "MES açısından kullanımı"],
        [
            [".../status", "ESP32", "mes_web", "Retained hat durumu, dashboard durum alanları"],
            [".../logs", "ESP32", "mes_web / workbook sink", "Olay logu, üretim tamamlanma ve ölçüm ayrıştırma"],
            [".../heartbeat", "ESP32 veya cihaz katmanı", "mes_web", "Cihaz canlılık kontrolü"],
            [".../bridge/status", "ESP32", "mes_web", "Wi-Fi, MQTT, queue ve drop telemetrisi"],
            [".../tablet/log", "Legacy tablet/audit akışı", "mes_web", "Fault, OEE snapshot ve audit satırları"],
            [".../cmd", "mes_web", "ESP32 -> Mega", "Start, stop, status, pickplace ve kalibrasyon komutları"],
            [".../vision/events", "Raspberry observer", "mes_web / workbook sink", "Pasif crossing ve renk çapraz kontrol kaydı"],
        ],
        [1.4, 1.15, 1.35, 2.4],
    )

    add_heading(doc, "4. Önerilen Çözüm Yöntemleri ve Gerekçeleri", 1)
    add_p(doc, "Önerilen çözüm yöntemi, üretim sistemini yedi katmanlı bir mimari olarak ele almaktır: fiziksel katman, kontrol katmanı, edge/bridge katmanı, haberleşme katmanı, MES uygulama katmanı, gözlem katmanı ve veri/entegrasyon katmanı. Bu yapı, siber-fiziksel sistem yaklaşımına uygundur; çünkü fiziksel hareket ve sensör verisi, siber tarafta çalışan veri işleme, dashboard, OEE ve karar destek mekanizmalarıyla sürekli ilişkilendirilmektedir.", after=6)
    add_p(doc, "Birinci gerekçe, gerçek zamanlı izlenebilirliktir. Operasyon sahasında ürün sayısı, renk kararı, robot çevrimi veya arıza bilgisi yalnızca seri monitörde kalırsa üretim yönetimi açısından kaybolur. `mes_web`, bu bilgileri canlı dashboard ve WebSocket snapshot'larına dönüştürerek operatör ve mühendis için görünür hale getirir. MES sistemlerinin önemi de burada ortaya çıkar: MES, sahadaki olayları yalnızca göstermekle kalmaz, onları iş emri, kalite, bakım, OEE ve raporlama bağlamına yerleştirir.", after=6)
    add_p(doc, "İkinci gerekçe, veri bütünlüğüdür. Excel workbook yapısı küçük ölçekli bir proje için pratik bir veri tabanı sınırı oluşturmaktadır. Kurumsal veri tabanına geçmeden önce olayların hangi sheet'e, hangi kolonlarla ve hangi kimliklerle yazıldığının netleşmesi gerekir. Bu yaklaşım, ileride FERP veya resmi raporlama katmanına geçiş için veri sözleşmesini olgunlaştırır.", after=6)
    add_p(doc, "Üçüncü gerekçe, OEE'nin doğru sınıflandırılmasıdır. Endüstri mühendisliği çalışmalarında OEE hesaplamak kadar, kaybın doğru kategoriye konması da önemlidir. Projede açılış kontrol listesi OEE dışı, kapanış/bakım süresi planlı duruş, manuel arıza süresi ise plansız duruş olarak işlenmektedir. ISO 22400'un MOM kapsamında KPI tanımlama ve kullanma çerçevesi, bu tür metriklerin sistematik biçimde ele alınması gerektiğini desteklemektedir (International Organization for Standardization, 2014a, 2014b).", after=6)
    add_p(doc, "Dördüncü gerekçe, düşük maliyetli erişilebilirliktir. Büyük ölçekli ticari MES sistemleri güçlü olmakla birlikte laboratuvar ve KOBİ ölçeğinde ilk kurulum için maliyetli ve karmaşık olabilir. Bu proje, açık kodlu web teknolojileri, MQTT, Arduino/ESP32 ve Excel tabanlı veri kayıt yapısı ile MES kavramlarının küçük ölçekte uygulanabileceğini göstermektedir. Bu yönüyle proje yalnızca teknik bir prototip değil, eğitim amaçlı uygulanabilir bir dijital üretim platformudur.", after=8)

    add_table(
        doc,
        "Tablo 3. Endüstri mühendisliği KPI ve veri eşleştirmesi",
        ["Karar alanı", "Projedeki veri", "MES çıktısı", "Yorum"],
        [
            ["Çevrim süresi", "PICKPLACE_DONE, RELEASED, flow_ms, cycle_ms", "Tamamlanan ürün ve OEE anlıkları", "Performans ve darboğaz analizi için temel"],
            ["Kalite", "Sensor rengi, vision rengi, quality override", "GOOD / REWORK / SCRAP kaydı", "Kalite kaybı ve düzeltme izlenebilirliği"],
            ["Duruş", "manualFaultDurationMs, activeFault", "3_Arizalar ve OEE runtime", "Plansız duruş etkisinin availability hesabına aktarılması"],
            ["Bakım", "maintenance checklist ve closing duration", "9_Bakim_Kayitlari", "Planlı duruş ve bakım disiplininin ayrıştırılması"],
            ["İş emri", "ordersById, activeOrderId, completionLog", "7_Is_Emirleri", "Üretim yürütme ve stok kontrol bağı"],
            ["Stok", "inventoryByProduct, SCRAP kuralı", "8_Depo_Stok", "Hurda ürünün stoka girmemesiyle veri tutarlılığı"],
        ],
        [1.15, 1.85, 1.65, 1.85],
    )

    add_heading(doc, "5. Beklenen Faydalar", 1)
    add_heading(doc, "A. Sanayiye Faydaları", 2)
    add_p(doc, "Sanayi açısından çalışmanın temel faydası, saha cihazlarından gelen olayların üretim yönetimi bağlamında anlamlandırılmasıdır. Bir işletmede yalnızca üretim adedi değil, bu adedin hangi iş emrinden, hangi kalite durumuyla, hangi duruş ve bakım koşulları altında üretildiği önemlidir. Projedeki MES katmanı bu bakış açısını küçük ölçekte ortaya koymaktadır.", after=6)
    add_p(doc, "İkinci sanayi faydası izlenebilirliktir. Ürün kimliği, ölçüm kimliği, renk kararı, robot çevrimi ve kalite düzeltmesi aynı veri akışında tutulduğunda, üretim sonrası hata analizi daha sistematik yapılabilir. Özellikle kalite override ve hurda ürünün stoka alınmaması gibi kurallar, üretim ve depo kayıtları arasındaki tutarsızlık riskini azaltır.", after=6)
    add_p(doc, "Üçüncü fayda bakım ve arıza yönetimidir. Operator kiosk üzerinden açılan arıza bildirimi teknisyen ekranına taşınmakta; `Cevapla` ve `Tamamla` aksiyonlarıyla cevap süresi, giderme süresi ve toplam süre takip edilebilmektedir. Bu yapı, ileride MTTR, çağrı cevap süresi ve bakım etkinliği gibi metriklerin geliştirilmesine zemin hazırlar.", after=6)
    add_p(doc, "Dördüncü fayda ölçeklenebilirliktir. Renk ayırma hattı basit görünse de mimaride kullanılan katmanlar daha büyük sistemlerde de karşılık bulur: saha kontrolü, haberleşme, MES, OEE, bakım, kalite ve raporlama. Bu nedenle proje ileride pick-to-light, vision doğrulama veya FERP entegrasyonu gibi modüllerle genişletilebilir.", after=6)

    add_heading(doc, "B. Akademiye ve Öğrencilere Faydaları", 2)
    add_p(doc, "Akademik açıdan çalışma, Endüstri 4.0, siber-fiziksel sistem, MES, OEE ve üretim veri modeli kavramlarını tek bir deney düzeneğinde birleştirmektedir. Öğrenciler yalnızca Arduino veya web arayüzü geliştirmeyi değil, bu teknolojilerin üretim yönetimi kararlarına nasıl bağlandığını da görebilmektedir.", after=6)
    add_p(doc, "İkinci akademik fayda, endüstri mühendisliği derslerinde soyut kalan kavramların canlı veriyle ilişkilendirilmesidir. OEE, kalite oranı, plansız duruş, planlı bakım, iş emri ve stok gibi kavramlar fiziksel hattaki olaylarla birlikte ele alındığında daha anlaşılır hale gelir.", after=6)
    add_p(doc, "Üçüncü akademik fayda, tekrar edilebilir araştırma platformudur. Aynı sistem farklı kalibrasyon koşulları, farklı iş emirleri, farklı arıza senaryoları ve farklı dashboard kararlarıyla yeniden çalıştırılabilir. Bu özellik, deneysel veri toplama ve karşılaştırmalı analiz için önemlidir.", after=8)

    add_heading(doc, "6. İlgili Literatür", 1)
    add_p(doc, "Endüstri 4.0 literatürü, üretim sistemlerinin birlikte çalışabilirlik, gerçek zamanlı veri, hizmet yönelimi ve modülerlik ilkeleriyle tasarlanması gerektiğini vurgular. Hermann, Pentek ve Otto (2016), Endüstri 4.0 senaryolarında bu ilkeleri temel tasarım prensipleri olarak ele almıştır. Projede kullanılan MQTT, REST, WebSocket ve modüler backend yapısı bu ilkelerle uyumludur.", after=6)
    add_p(doc, "Siber-fiziksel sistemler literatürü, fiziksel üretim ortamı ile siber hesaplama alanının veriyle birbirine bağlanmasını merkeze alır. Lee et al. (2015), CPS tabanlı üretim sistemlerinde fiziksel varlıkların durumunun siber alanda izlenmesi ve analiz edilmesini temel gereksinimlerden biri olarak görür. Bu projede fiziksel sensör ve robot olaylarının `mes_web` snapshot'larına dönüşmesi bu ilişkinin laboratuvar ölçeğindeki karşılığıdır.", after=6)
    add_p(doc, "MES ve ISA-95 literatürü, üretim yürütme sistemlerinin işletme planlama katmanı ile saha kontrol katmanı arasındaki rolünü açıklar. Govindaraju, Lukman ve Chandra (2014), MES tasarımında ISA-95 kullanımını ele alırken; ISA-95 standardı MOM/MES fonksiyonlarını seviye 3 bağlamında konumlandırır. Bu proje de Mega ve ESP32 ile saha katmanını, `mes_web` ile MES katmanını, workbook/FERP hazırlığı ile üst seviye raporlama katmanını ayrıştırmaktadır.", after=6)
    add_p(doc, "OEE literatürü, üretim etkinliğinin kullanılabilirlik, performans ve kalite bileşenleri üzerinden değerlendirilmesini sağlar. Nakajima (1988) OEE yaklaşımını TPM bağlamında yaygınlaştırırken, Ng Corrales et al. (2020) OEE'nin farklı sektör ve yaklaşımlarda gelişen bir KPI olduğunu göstermektedir. Projede OEE, aktif vardiya, hedef miktar, ideal çevrim, planlı duruş ve plansız arıza süreleriyle birlikte hesaplanmaktadır.", after=6)
    add_p(doc, "Dijital ikiz literatürü, fiziksel sistemin sanal karşılığı ile veri bağlantısını ele alır. Kritzinger et al. (2018), dijital model, dijital gölge ve dijital ikiz kavramlarını veri entegrasyon düzeyine göre ayırır. Projedeki mevcut yapı tam kapsamlı dijital ikiz değildir; ancak fiziksel olayların canlı dijital temsilini üreterek dijital gölge ve ileride dijital ikiz geliştirme için gerekli veri zeminini oluşturmaktadır.", after=8)

    doc.add_page_break()
    add_table(
        doc,
        "Tablo 4. Sistem işleyişi, kontrol noktası ve rapor çıktısı",
        ["Saha adımı", "Kontrol noktası", "Olası risk", "MES / rapor çıktısı"],
        [
            ["Ürün algılama", "Renk sensörü ve objectPresent", "Işık değişimi veya yanlış okuma", "Ölçüm logu, review_required potansiyeli"],
            ["Renk kararı", "Medyan, oy ve skor tabanlı sınıflandırma", "Kırmızı-sarı ayrımında kararsızlık", "TCS3200 ölçüm satırı ve karar kaynağı"],
            ["Kuyruğa alma", "pendingCount ve item_id", "Kuyruk dolması veya takip kaybı", "queue_enq olayı ve ürün kimliği"],
            ["Robot alma-bırakma", "LIM22/LIM23 ve timeout", "Sıkışma veya hedefe ulaşamama", "pickplace eventleri ve arıza nedeni"],
            ["Operator müdahalesi", "Kiosk aksiyonları", "Kayıtsız duruş veya kalite düzeltmesi", "Fault, help request, quality override"],
            ["Raporlama", "Workbook flush ve sheet düzeni", "Veri kaybı veya yanlış sınıflandırma", "Günlük Excel, OEE runtime ve audit izi"],
        ],
        [1.3, 1.6, 1.45, 1.85],
    )

    add_heading(doc, "7. Hedeflenen Çıktılar", 1)
    add_p(doc, "Çalışmanın birinci hedef çıktısı, MES tabanlı siber-fiziksel konveyör sistemi için akademik ara rapordur. Bu rapor, fiziksel sistem açıklaması ile MES, CPS, MQTT, OEE ve veri tabanı katmanlarını endüstri mühendisliği bakışıyla birleştirmektedir.", after=6)
    add_p(doc, "İkinci hedef çıktı, fiziksel sistem ve haberleşme mimarisinin açık biçimde dokümante edilmesidir. Mega, ESP32, MQTT broker, `mes_web`, dashboard, kiosk, teknisyen ekranı ve Excel workbook arasındaki veri akışı raporda tablo ve şekil önerileriyle açıklanmıştır.", after=6)
    add_p(doc, "Üçüncü hedef çıktı, OEE ve üretim izlenebilirliği için ölçülebilir veri modelinin tanımlanmasıdır. Tamamlanan ürün, kalite durumu, duruş türü, bakım süresi, iş emri ve stok davranışları aynı MES veri katmanında izlenebilir hale getirilmektedir.", after=6)
    add_p(doc, "Dördüncü hedef çıktı, ileriki geliştirmeler için akademik ve teknik yol haritasıdır. Vision çapraz kontrolünün sahada doğrulanması, workbook replay aracı, FERP JSON kontratı, yerel broker seçeneği ve çok istasyonlu genişleme ileriki çalışma başlıkları olarak değerlendirilmektedir.", after=8)

    doc.add_page_break()
    add_table(
        doc,
        "Tablo 5. Kullanılan literatür ve bu çalışmadaki rolü",
        ["Konu", "Temel kaynak", "Bu çalışmadaki kullanımı"],
        [
            ["Endüstri 4.0", "Hermann et al. (2016)", "Gerçek zamanlılık, modülerlik ve birlikte çalışabilirlik gerekçesi"],
            ["CPS", "Lee et al. (2015)", "Fiziksel hat ile siber MES katmanı arasındaki veri bağlantısı"],
            ["MES / ISA-95", "ISA-95; Govindaraju et al. (2014); Kletti (2007)", "MES'in ERP-saha kontrol köprüsü olarak konumlandırılması"],
            ["MQTT", "OASIS (2014)", "Topic tabanlı IoT haberleşmesinin teknik dayanağı"],
            ["OEE / KPI", "Nakajima (1988); ISO 22400; Ng Corrales et al. (2020)", "Availability, performance ve quality metriklerinin yorumlanması"],
            ["Dijital ikiz", "Kritzinger et al. (2018); Lu et al. (2020)", "Canlı veri, dijital gölge ve ileride dijital ikiz geliştirme temeli"],
        ],
        [1.1, 2.05, 3.1],
    )

    add_heading(doc, "Önerilen Şekil ve Tablo Listesi", 2)
    add_bullet(doc, "Şekil 1: Fiziksel konveyör, robot kol, renk sensörü ve limit switch konumlarını gösteren saha fotoğrafı.")
    add_bullet(doc, "Şekil 2: Mega, ESP32, MQTT broker, mes_web, dashboard, kiosk, teknisyen ekranı ve Excel workbook veri akışı.")
    add_bullet(doc, "Şekil 3: MQTT topic ağacı ve publish/subscribe ilişkisi.")
    add_bullet(doc, "Şekil 4: OEE bileşenleri ile planlı duruş, plansız duruş ve kalite düzeltme ilişkisi.")
    add_bullet(doc, "Şekil 5: Workbook sheet yapısı ve FERP/dijital ikiz genişleme noktaları.")
    add_bullet(doc, "Tablo 1: MES-CPS mimari katman haritası.")
    add_bullet(doc, "Tablo 2: MQTT topic ağacı ve veri akışı özeti.")
    add_bullet(doc, "Tablo 3: Endüstri mühendisliği KPI ve veri eşleştirmesi.")
    add_bullet(doc, "Tablo 4: Sistem işleyişi, kontrol noktası ve rapor çıktısı.")
    add_bullet(doc, "Tablo 5: Kullanılan literatür ve bu çalışmadaki rolü.")

    add_heading(doc, "8. Kullanılan Güncel Literatür Örnekleri", 1)
    add_p(doc, "Bu bölümdeki kaynaklar, çalışmanın üretim yürütme, siber-fiziksel sistem, IoT haberleşme, OEE/KPI ve dijital ikiz boyutlarını desteklemek için seçilmiştir. Raporun amacı ticari MES ürünü önermek değil, akademik ve uygulanabilir bir üretim veri mimarisi kurmaktır. Bu nedenle kaynaklar standartlar, hakemli çalışmalar ve temel kitaplar etrafında sınırlandırılmıştır.", after=6)
    add_p(doc, "ISA-95 ve ISO 22400 kaynakları, sistemin işletme-kontrol ayrımı ve KPI yönetimi açısından kavramsal temelini oluşturur. MQTT standardı, Mega-ESP32-MES arasındaki topic tabanlı haberleşmenin teknik arka planını açıklar. OEE ve dijital ikiz literatürü ise endüstri mühendisliği metriklerinin canlı veri ve simülasyon potansiyeli ile nasıl ilişkilendirilebileceğini göstermektedir.", after=8)

    add_heading(doc, "Sonuç Niteliğinde Ara Değerlendirme", 1)
    add_p(doc, "Bu ara rapor, mevcut mini konveyör projesinin yalnızca mekanik ve elektronik bir otomasyon sistemi olmadığını; MES, OEE, izlenebilirlik, arıza/bakım yönetimi ve siber-fiziksel üretim yaklaşımı açısından değerlendirilebilecek bütünleşik bir çalışma olduğunu göstermektedir. Arduino Mega fiziksel karar otoritesi, ESP32 haberleşme köprüsü, MQTT veri taşıma katmanı, `mes_web` ise üretim yürütme ve raporlama katmanı olarak konumlandırılmıştır.", after=6)
    add_p(doc, "Çalışmanın en güçlü yönü, düşük maliyetli bileşenlerle endüstriyel kavramların uygulanabilir hale getirilmesidir. Dashboard, kiosk, teknisyen ekranı, OEE runtime ve günlük workbook yapısı; öğrencilerin ve araştırmacıların üretim olaylarını yalnızca izlemekle kalmayıp analiz edebilmesine de imkan vermektedir.", after=6)
    add_p(doc, "Çalışmanın sınırlılığı, deney ortamının laboratuvar ölçeğinde olması ve tüm endüstriyel değişkenleri kapsamamasıdır. Gerçek fabrika ortamında çoklu istasyon, güvenlik sertifikasyonu, kullanıcı yetkilendirme, veri tabanı dayanıklılığı, ağ güvenliği ve ERP entegrasyonu gibi ek gereksinimler ortaya çıkacaktır. Buna rağmen mevcut mimari, saha verisinin MES mantığıyla nasıl yapılandırılabileceğini göstermesi bakımından güçlü bir ara rapor temelidir.", after=6)
    add_p(doc, "Gelecek çalışmalarda fiziksel sistem fotoğrafı ve bağlantı şeması eklenmeli, daha uzun süreli üretim testleri yapılmalı, kontrollü arıza senaryoları çalıştırılmalı, vision çapraz kontrolü sahada doğrulanmalı ve workbook verisinden FERP/dijital ikiz katmanına aktarılacak resmi veri kontratı netleştirilmelidir.", after=8)

    add_heading(doc, "Proje İçi Kaynak Dosyalar", 1)
    add_bullet(doc, "README/architecture.md - MES, fiziksel katman, kontrol katmanı, edge/bridge, haberleşme, uygulama ve veri katmanı mimari açıklamaları.")
    add_bullet(doc, "README/hardware.md - Mega, ESP32, renk sensörü, limit switch, konveyör, robot kol ve operator bilgisayarı rol dağılımı.")
    add_bullet(doc, "README/mqtt-topics.md - Aktif MQTT topic ailesi, publisher/subscriber rolleri ve komut payload notları.")
    add_bullet(doc, "README/data-model.md - Dashboard snapshot, kiosk snapshot, teknisyen snapshot, OEE runtime ve workbook sheet modeli.")
    add_bullet(doc, "mes_web/README.md - FastAPI, REST, WebSocket, kiosk, teknisyen, OEE ve workbook kayıt davranışı.")
    add_bullet(doc, "CPP/mega.cpp - Fiziksel karar, renk sınıflandırma, robot kol, limit switch, kuyruk ve status/log üretim kodu.")
    add_bullet(doc, "CPP/esp32.cpp - Mega seri hattı ile MQTT arasında bridge görevi gören ESP32 kodu.")
    add_bullet(doc, "raspberry/README.md - Pasif vision observer, crossing ve renk çapraz kontrol akışı.")

    add_heading(doc, "Kaynakça", 1)
    refs = [
        "Govindaraju, R., Lukman, K., & Chandra, D. R. (2014). Manufacturing Execution System Design Using ISA-95. Advanced Materials Research, 980, 248-252. https://doi.org/10.4028/www.scientific.net/AMR.980.248",
        "Hermann, M., Pentek, T., & Otto, B. (2016). Design principles for Industrie 4.0 scenarios. 2016 49th Hawaii International Conference on System Sciences (HICSS), 3928-3937. https://doi.org/10.1109/HICSS.2016.488",
        "International Organization for Standardization. (2014a). ISO 22400-1:2014 Automation systems and integration - Key performance indicators (KPIs) for manufacturing operations management - Part 1: Overview, concepts and terminology. https://www.iso.org/standard/56847.html",
        "International Organization for Standardization. (2014b). ISO 22400-2:2014 Automation systems and integration - Key performance indicators (KPIs) for manufacturing operations management - Part 2: Definitions and descriptions. https://www.iso.org/standard/54497.html",
        "International Society of Automation. (n.d.). ISA-95 standard: Enterprise-control system integration. Retrieved April 25, 2026, from https://www.isa.org/standards-and-publications/isa-standards/isa-95-standard",
        "Kletti, J. (Ed.). (2007). Manufacturing Execution Systems - MES. Springer. https://doi.org/10.1007/978-3-540-49744-8",
        "Kritzinger, W., Karner, M., Traar, G., Henjes, J., & Sihn, W. (2018). Digital Twin in manufacturing: A categorical literature review and classification. IFAC-PapersOnLine, 51(11), 1016-1022. https://doi.org/10.1016/j.ifacol.2018.08.474",
        "Lee, J., Bagheri, B., & Kao, H.-A. (2015). A cyber-physical systems architecture for Industry 4.0-based manufacturing systems. Manufacturing Letters, 3, 18-23. https://doi.org/10.1016/j.mfglet.2014.12.001",
        "Lu, Y., Liu, C., Wang, K. I.-K., Huang, H., & Xu, X. (2020). Digital Twin-driven smart manufacturing: Connotation, reference model, applications and research issues. Robotics and Computer-Integrated Manufacturing, 61, Article 101837. https://doi.org/10.1016/j.rcim.2019.101837",
        "Nakajima, S. (1988). Introduction to TPM: Total Productive Maintenance. Productivity Press.",
        "Ng Corrales, L. C., Lambán, M. P., Hernandez Korner, M. E., & Royo, J. (2020). Overall equipment effectiveness: Systematic literature review and overview of different approaches. Applied Sciences, 10(18), Article 6469. https://doi.org/10.3390/app10186469",
        "OASIS. (2014). MQTT Version 3.1.1. OASIS Standard. https://docs.oasis-open.org/mqtt/mqtt/v3.1.1/mqtt-v3.1.1.html",
    ]
    add_references(doc, refs)

    core = doc.core_properties
    core.title = "MES Tabanlı Siber-Fiziksel Konveyör Sistemi Ara Rapor Taslağı"
    core.subject = "MES, OEE, CPS, MQTT ve konveyör ara raporu"
    core.keywords = "MES, CPS, OEE, MQTT, Konveyör, Endüstri Mühendisliği"
    core.comments = "Generated from project README and code context."

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
